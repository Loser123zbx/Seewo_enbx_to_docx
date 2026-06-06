# -*- coding: utf-8 -*-
import os
import sys
import json


with open(os.path.join(os.path.dirname(__file__), 'options.json'), 'r', encoding='utf-8') as f:
        options = json.load(f)
        text_size_proportion = options.get("text_size_proportion")


# 确保可以导入同目录下的 read_xml
sys.path.append(os.path.dirname(__file__))

def hex_to_rgb(hex_color):
        """
        将 HEX 颜色字符串 (#RRGGBB 或 #AARRGGBB) 转换为 RGB 元组 (R, G, B)
        """
        if not hex_color:
                return (0, 0, 0)
        
        # 去除 # 号
        hex_color = hex_color.lstrip('#')
        
        # 处理带透明度的颜色 (#AARRGGBB)，只取后6位 RRGGBB
        if len(hex_color) == 8:
                hex_color = hex_color[2:]
        elif len(hex_color) != 6:
                return (0, 0, 0)
                
        try:
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                return (r, g, b)
        except ValueError:
                return (0, 0, 0)

def convert_xml_to_docx(structured_data, output_path=None, doc=None):
        """
        将 read_xml_to_sentence 返回的结构化数据转换为 docx 文件
        
        参数:
        structured_data: list[list[list[TextRunData]]] 
                         结构: [ TextBox [ Line [ Run ] ] ]
        output_path: str 输出文件路径
        """
        # 只在需要时导入 docx 模块
        from docx import Document
        from docx.shared import Pt, RGBColor
        from read_xml import TextRunData

        created_doc = False
        if doc is None:
                doc = Document()
                created_doc = True
        
        # 遍历每个文本框
        for textbox_index, textbox_lines in enumerate(structured_data):
                print(f"正在处理文本框 {textbox_index+1}...{textbox_lines}")
                # 如果不是第一个文本框，添加一个空行作为间隔（可选）
                if textbox_index > 0:
                        doc.add_paragraph()
                
                # 遍历文本框中的每一行
                for line_runs in textbox_lines:
                        # 创建新段落
                        paragraph = doc.add_paragraph()
                        
                        # 如果行为空，跳过
                        if not line_runs:
                                continue
                        
                        # 遍历行中的每个文本片段 (Run)
                        for run_data in line_runs:
                                text_content = run_data.text
                                
                                # 跳过纯空白或换行符（Word段落自动换行）
                                # 注意：如果希望保留行内换行，需特殊处理，这里假设每行对应一个段落
                                if not text_content or text_content.strip() == '':
                                        # 如果是空串但可能有格式占位需求，可以添加空run，通常不需要
                                        continue
                                
                                # 添加 Run
                                run = paragraph.add_run(text_content)
                                
                                # --- 应用格式 ---
                                
                                # 1. 字体大小
                                if run_data.font_size > 0:
                                        from math import ceil
                                        run.font.size = ceil(Pt(run_data.font_size)*float(text_size_proportion))
                                
                                # 2. 加粗
                                if run_data.is_bold:
                                        run.font.bold = True
                                
                                # 3. 斜体
                                if run_data.is_italic:
                                        run.font.italic = True
                                
                                # 4. 字体名称
                                if run_data.font_family:
                                        run.font.name = run_data.font_family
                                
                                # 5. 字体颜色 (前景色)
                                if run_data.foreground_color and run_data.foreground_color != "#FF000000":
                                        rgb = hex_to_rgb(run_data.foreground_color)
                                        run.font.color.rgb = RGBColor(*rgb)
                                        
                                # 6. 背景色 (高亮) - docx 对背景色支持有限，通常用 Highlight
                                # 注意：Word 的 Highlight 颜色选项有限，这里仅做简单映射或忽略
                                # 如果需要精确背景色，通常需要更复杂的 shading 设置，此处暂略以保持简洁

        # 保存文档（如果 caller 指定了 output_path 或我们自己创建了 doc）
        if created_doc and output_path:
                try:
                        doc.save(output_path)
                        print(f"文档已保存至: {output_path}")
                except Exception as e:
                        print(f"保存文档失败: {e}")
                        raise

def convert_xml_to_docx_English2ChineseMode(structured_data, output_path=None, doc=None):
        """
        将 read_xml_to_sentence 返回的结构化数据转换为 docx 文件
        特殊逻辑：如果只有两个文本框，则将对应行合并为中英对照
        """
        # 只在需要时导入 docx 模块
        from docx import Document
        from docx.shared import Pt, RGBColor
        from read_xml import TextRunData
        from math import ceil

        created_doc = False
        if doc is None:
                doc = Document()
                created_doc = True

        # 支持 structured_data 可能来自单个 slide（2 个文本框）或多个 slide 的情况。
        # 我们按 slide 单位处理：如果 structured_data 看起来是多个文本框集合（len>2），
        # 则逐个尝试按 2 个文本框为一组处理；否则直接按单个 slide 处理。

        def _process_pair(en_box, cn_box):
                max_lines = max(len(en_box), len(cn_box))
                for i in range(max_lines):
                        paragraph = doc.add_paragraph()
                        if i < len(en_box):
                                for run_data in en_box[i]:
                                        try:
                                                text_content = run_data.text
                                        except Exception:
                                                text_content = ''
                                        if not text_content or text_content.strip() == '':
                                                continue
                                        run = paragraph.add_run(text_content)
                                        if run_data.font_size > 0:
                                                run.font.size = ceil(Pt(run_data.font_size)*float(text_size_proportion))
                                        if run_data.is_bold:
                                                run.font.bold = True
                                        if run_data.is_italic:
                                                run.font.italic = True
                                        if run_data.font_family:
                                                run.font.name = run_data.font_family
                                        if run_data.foreground_color and run_data.foreground_color != "#FF000000":
                                                rgb = hex_to_rgb(run_data.foreground_color)
                                                run.font.color.rgb = RGBColor(*rgb)
                                paragraph.add_run("    ")
                        if i < len(cn_box):
                                for run_data in cn_box[i]:
                                        try:
                                                text_content = run_data.text
                                        except Exception:
                                                text_content = ''
                                        if not text_content or text_content.strip() == '':
                                                continue
                                        run = paragraph.add_run(text_content)
                                        if run_data.font_size > 0:
                                                run.font.size = ceil(Pt(run_data.font_size)*float(text_size_proportion))
                                        if run_data.is_bold:
                                                run.font.bold = True
                                        if run_data.is_italic:
                                                run.font.italic = True
                                        if run_data.font_family:
                                                run.font.name = run_data.font_family
                                        if run_data.foreground_color and run_data.foreground_color != "#FF000000":
                                                rgb = hex_to_rgb(run_data.foreground_color)
                                                run.font.color.rgb = RGBColor(*rgb)

        # If structured_data is a list of textboxes (possibly many), process accordingly
        if isinstance(structured_data, list) and len(structured_data) > 0 and all(isinstance(x, list) for x in structured_data):
                if len(structured_data) == 2 and all(isinstance(x[0] if x else [], list) for x in structured_data):
                        en_box = structured_data[0]
                        cn_box = structured_data[1]
                        _process_pair(en_box, cn_box)
                else:
                        i = 0
                        n = len(structured_data)
                        while i < n:
                                if i+1 < n:
                                        en_box = structured_data[i]
                                        cn_box = structured_data[i+1]
                                        _process_pair(en_box, cn_box)
                                        i += 2
                                else:
                                        # leftover single textbox, render normally
                                        for line_runs in structured_data[i]:
                                                p = doc.add_paragraph()
                                                for run_data in line_runs:
                                                        if not run_data.text or run_data.text.strip() == '':
                                                                continue
                                                        r = p.add_run(run_data.text)
                                                        if run_data.font_size > 0:
                                                                r.font.size = ceil(Pt(run_data.font_size)*float(text_size_proportion))
                                                        if run_data.is_bold:
                                                                r.font.bold = True
                                                        if run_data.is_italic:
                                                                r.font.italic = True
                                                        if run_data.font_family:
                                                                r.font.name = run_data.font_family
                                                        if run_data.foreground_color and run_data.foreground_color != "#FF000000":
                                                                rgb = hex_to_rgb(run_data.foreground_color)
                                                                r.font.color.rgb = RGBColor(*rgb)
                                        i += 1

        # 保存文档
        if created_doc and output_path:
                try:
                        doc.save(output_path)
                        print(f"文档已保存至: {output_path}")
                except Exception as e:
                        print(f"保存文档失败: {e}")
                        raise

if __name__ == '__main__':
        # 1. 导入解析函数
        from read_xml import read_xml_to_sentence
        
        # 2. 解析 XML
        xml_path = 'D:\\123_softs\\enbx_to_docx\\src\\core\\tmp\\Slides\\Slide_47.xml'
        output_docx = 'D:\\123_softs\\enbx_to_docx\\src\\core\\tmp\\Slide_47.docx'
        
        print(f"正在解析: {xml_path}")
        structured_data = read_xml_to_sentence(xml_path)
        
        if structured_data:
                print(f"解析成功，共 {len(structured_data)} 个文本框。正在生成 DOCX...")
                # 3. 转换为 DOCX (使用中英对照模式)
                convert_xml_to_docx_English2ChineseMode(structured_data, output_docx)
        else:
                print("解析失败或无数据")