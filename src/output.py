# -*- coding: utf-8 -*-

#--------------------------------------------------------------------------
# Python code generated with wxFormBuilder (version 3.9.0 Jun 14 2020)
# http://www.wxformbuilder.org/
#
# PLEASE DO *NOT* EDIT THIS FILE!
#--------------------------------------------------------------------------

import wx
import os
import wx.xrc
import re
import gettext
_ = gettext.gettext

#--------------------------------------------------------------------------
#  Class output
#---------------------------------------------------------------------------

class output ( wx.Frame ):

        def __init__(self, parent):
                import os
                wx.Frame.__init__ (self, parent, id = wx.ID_ANY, title = wx.EmptyString, pos = wx.DefaultPosition, size = wx.Size( 800,600 ), style = wx.DEFAULT_FRAME_STYLE|wx.TAB_TRAVERSAL)

                self.SetSizeHints(wx.DefaultSize, wx.DefaultSize)
                self.SetBackgroundColour(wx.Colour( 239, 235, 235 ))

                root = wx.BoxSizer(wx.HORIZONTAL)


                # 加载图片 "output_setting.png"
                img_path = None
                # 尝试多个可能的路径
                possible_paths = [
                        os.path.join(os.path.dirname(__file__), 'output_setting.png'),
                        os.path.join(os.path.dirname(__file__), 'src', 'output_setting.png'),
                        'src/output_setting.png',
                        'output_setting.png'
                ]
                for p in possible_paths:
                        if os.path.isfile(p):
                                img_path = p
                                break
                if img_path:
                        try:
                                img = wx.Image(img_path, wx.BITMAP_TYPE_ANY)
                                # 限制图片最大宽度，防止过大
                                max_width = 1200
                                if img.GetWidth() > max_width:
                                        scale = max_width / float(img.GetWidth())
                                        new_height = int(img.GetHeight() * scale)
                                        img = img.Scale(max_width, new_height, wx.IMAGE_QUALITY_HIGH)
                                bitmap = wx.Bitmap(img)
                                self.image_ctrl = wx.StaticBitmap(self, wx.ID_ANY, bitmap, wx.DefaultPosition, wx.DefaultSize, 0)
                                root.Add(self.image_ctrl, 0, wx.ALL|wx.EXPAND, 5)
                        except Exception as e:
                                print(f"加载图片失败: {e}")

                settings = wx.BoxSizer(wx.VERTICAL)

                self.output_text = wx.StaticText(self, wx.ID_ANY, _(u"填充导出选项"), wx.DefaultPosition, wx.DefaultSize, 0)
                self.output_text.Wrap(-1)
                self.output_text.SetFont(wx.Font( 15, wx.FONTFAMILY_DEFAULT, wx.FONTSTYLE_NORMAL, wx.FONTWEIGHT_NORMAL, False, wx.EmptyString ))
                settings.Add(self.output_text, 0, wx.ALL, 5)

                self.slide_list_label = wx.StaticText(self, wx.ID_ANY, _(u"选择要导出的版面 (Slides):"), wx.DefaultPosition, wx.DefaultSize, 0)
                settings.Add(self.slide_list_label, 0, wx.ALL, 5)
                # 顶部添加全选按钮
                h_sel = wx.BoxSizer(wx.HORIZONTAL)
                self.select_all_btn = wx.Button(self, wx.ID_ANY, _(u"全选/取消全选"), wx.DefaultPosition, wx.DefaultSize, 0)
                h_sel.Add(self.select_all_btn, 0, wx.ALL, 3)
                # 添加“选项”按钮在同一行
                self.options_btn = wx.Button(self, wx.ID_ANY, _(u"选项"), wx.DefaultPosition, wx.DefaultSize, 0)
                h_sel.Add(self.options_btn, 0, wx.ALL, 3)
                settings.Add(h_sel, 0, wx.EXPAND, 0)
                self.slide_list = wx.CheckListBox(self, wx.ID_ANY, wx.DefaultPosition, wx.Size(-1,120), [], wx.LB_ALWAYS_SB)
                settings.Add(self.slide_list, 0, wx.ALL|wx.EXPAND, 5)
                self.select_all_btn.Bind(wx.EVT_BUTTON, self.on_select_all)
                self.options_btn.Bind(wx.EVT_BUTTON, self.on_options)

                self.out_path_label = wx.StaticText(self, wx.ID_ANY, _(u"导出目录:"), wx.DefaultPosition, wx.DefaultSize, 0)
                settings.Add(self.out_path_label, 0, wx.ALL, 5)
                self.out_dirpicker = wx.DirPickerCtrl(self, wx.ID_ANY, os.path.abspath('.'), _(u"选择导出目录"), wx.DefaultPosition, wx.DefaultSize, wx.DIRP_DEFAULT_STYLE)
                settings.Add(self.out_dirpicker, 0, wx.ALL|wx.EXPAND, 5)

                self.fontsize = wx.StaticText(self, wx.ID_ANY, _(u"导出字符大小比例(从希沃课件到.docx文档的文字大小比例)"), wx.DefaultPosition, wx.DefaultSize, 0)
                self.fontsize.Wrap(-1)
                settings.Add(self.fontsize, 0, wx.ALL, 5)

                self.font_size_bili = wx.TextCtrl(self, wx.ID_ANY, _(u"0.3"), wx.DefaultPosition, wx.DefaultSize, 0)
                settings.Add(self.font_size_bili, 0, wx.ALL, 5)

                # 导出文件名模板
                self.filename_label = wx.StaticText(self, wx.ID_ANY, _(u"导出文件名模板（使用{slide}占位符）:"), wx.DefaultPosition, wx.DefaultSize, 0)
                settings.Add(self.filename_label, 0, wx.ALL, 5)
                self.filename_tpl = wx.TextCtrl(self, wx.ID_ANY, _(u"{slide}.docx"), wx.DefaultPosition, wx.DefaultSize, 0)
                settings.Add(self.filename_tpl, 0, wx.ALL|wx.EXPAND, 5)

                output_settingsChoices = [_(u"使用英汉互译模式导出"), _(u"只导出文本而不导出字符格式"), _(u"导出为同一文件")]

                self.output_settings = wx.CheckListBox(self, wx.ID_ANY, wx.DefaultPosition, wx.DefaultSize, output_settingsChoices, wx.LB_ALWAYS_SB)
                settings.Add(self.output_settings, 0, wx.ALL|wx.EXPAND, 5)
                self.output_settings.Check(2, True)  # 默认选中第一个选项

                self.output = wx.Button(self, wx.ID_ANY, _(u"导出"), wx.DefaultPosition, wx.DefaultSize, 0)
                settings.Add(self.output, 0, wx.ALL, 5)

                self.close = wx.Button(self, wx.ID_ANY, _(u"关闭"), wx.DefaultPosition, wx.DefaultSize, 0)
                settings.Add(self.close, 0, wx.ALL, 5)

                root.Add(settings, 1, wx.EXPAND, 5)


                root.Add(settings, 1, wx.EXPAND, 5)


                self.SetSizer( root )
                self.Layout()

                self.Centre(wx.BOTH)
        
                self.output.Bind(wx.EVT_BUTTON, self.on_export)
                self.close.Bind(wx.EVT_BUTTON, self.on_close)
                self.EnableCloseButton(False)     

        def on_close(self, event):
                self.Hide()
        def populate_slides(self, slides_path ="../tmp/Slides"):
                import os
                """从解压后的 Slides 目录填充可选版面列表（只显示文件名）"""
                # 规范化路径
                slides_path = os.path.normpath(slides_path)
                
                if not slides_path or not os.path.isdir(slides_path):
                        return
                # 按界面序号排序（Slide_数字.xml），否则按文件名
                files = [f for f in os.listdir(slides_path) if f.lower().endswith('.xml')]
                def slide_key(name):
                        m = re.search(r"(\d+)", name)
                        return int(m.group(1)) if m else 10**9
                files = sorted(files, key=slide_key)
                self.slide_list.Clear()
                for f in files:
                        self.slide_list.Append(f)

        def get_selected_slides(self):
                checked = self.slide_list.GetCheckedStrings()
                return list(checked)

        def on_select_all(self, event):
                # 如果有未选中的则全部选中，否则全部取消
                total = self.slide_list.GetCount()
                if total == 0:
                        return
                checked_count = len(self.slide_list.GetCheckedStrings())
                if checked_count < total:
                        # select all
                        for i in range(total):
                                self.slide_list.Check(i, True)
                else:
                        for i in range(total):
                                self.slide_list.Check(i, False)

        def on_options(self, event):
                import json
                opts_path = os.path.join(os.path.dirname(__file__), 'options.json')
                # 读取现有选项
                opts = {}
                try:
                        with open(opts_path, 'r', encoding='utf-8') as f:
                                opts = json.load(f)
                except Exception:
                        opts = {}

                dlg = wx.Dialog(self, title=_(u"导出选项"), size=(360,180))
                sizer = wx.BoxSizer(wx.VERTICAL)
                lbl = wx.StaticText(dlg, label=_(u"导出字符大小比例(text_size_proportion):"))
                sizer.Add(lbl, 0, wx.ALL, 8)
                val = str(opts.get('text_size_proportion', 0.3))
                self.opt_text_bili = wx.TextCtrl(dlg, value=val)
                sizer.Add(self.opt_text_bili, 0, wx.ALL|wx.EXPAND, 8)

                btns = wx.BoxSizer(wx.HORIZONTAL)
                save = wx.Button(dlg, label=_(u"保存"))
                cancel = wx.Button(dlg, label=_(u"取消"))
                btns.Add(save, 0, wx.ALL, 8)
                btns.Add(cancel, 0, wx.ALL, 8)
                sizer.Add(btns, 0, wx.ALIGN_CENTER)
                dlg.SetSizer(sizer)

                def on_save(evt):
                        try:
                                new_bili = float(self.opt_text_bili.GetValue())
                        except Exception:
                                wx.MessageBox(_(u"请输入有效的小数值"), _(u"错误"), wx.OK|wx.ICON_ERROR)
                                return
                        opts['text_size_proportion'] = new_bili
                        with open(opts_path, 'w', encoding='utf-8') as f:
                                json.dump(opts, f, ensure_ascii=False, indent=2)
                        # 立即更新已加载模块的变量
                        try:
                                import importlib
                                import text_to_docx
                                importlib.reload(text_to_docx)
                        except Exception:
                                pass
                        dlg.EndModal(wx.ID_OK)

                def on_cancel(evt):
                        dlg.EndModal(wx.ID_CANCEL)

                save.Bind(wx.EVT_BUTTON, on_save)
                cancel.Bind(wx.EVT_BUTTON, on_cancel)
                dlg.ShowModal()
                dlg.Destroy()

        def on_export(self, event):
                # 收集参数
                import os
                from docx import Document
                from docx.shared import Pt, RGBColor
                from math import ceil
                from read_xml import read_xml_to_sentence
                import text_to_docx
                
                out_dir = self.out_dirpicker.GetPath()
                if not out_dir:
                        wx.MessageBox(_(u"请选择导出目录。"), _(u"错误"), wx.OK|wx.ICON_ERROR)
                        return

                try:
                        bili = float(self.font_size_bili.GetValue())
                except Exception:
                        bili = 0.3

                choices = [self.output_settings.GetString(i) for i in range(self.output_settings.GetCount())]
                checked = [self.output_settings.IsChecked(i) for i in range(self.output_settings.GetCount())]
                use_en_cn_mode = choices[0] in choices and checked[0] if len(choices)>0 else False
                text_only = choices[1] in choices and checked[1] if len(choices)>1 else False
                same_file = choices[2] in choices and checked[2] if len(choices)>2 else False

                selected = self.get_selected_slides()
                if not selected:
                        # 如果没有选择，默认导出所有
                        selected = list(self.slide_list.GetItems())

                if not selected:
                        wx.MessageBox(_(u"没有可导出的幻灯片。"), _(u"提示"), wx.OK|wx.ICON_INFORMATION)
                        return

                # 定义源文件目录
                src_base_dir = "../tmp/Slides"

                errors = []
                combined_structured = [] if same_file else None
                total = len(selected)
                
                # 同步执行导出，不使用多线程
                for idx, slide in enumerate(selected):
                        src = os.path.join(src_base_dir, slide)
                        try:
                                structured = read_xml_to_sentence(src)
                        except Exception as e:
                                errors.append((slide, str(e)))
                                continue

                        # 决定输出路径/文件名
                        if same_file:
                                # 将解析结果合并
                                combined_structured.extend(structured)
                        else:
                                tpl = self.filename_tpl.GetValue() or "{slide}.docx"
                                name = tpl.replace('{slide}', os.path.splitext(slide)[0])
                                out_file = os.path.join(out_dir, name)

                                try:
                                        if text_only:
                                                # 简单文本导出
                                                d = Document()
                                                for textbox in structured:
                                                        for line_runs in textbox:
                                                                line_text = ''.join([r.text for r in line_runs if r.text])
                                                                if line_text.strip():
                                                                        d.add_paragraph(line_text)
                                                d.save(out_file)
                                        else:
                                                if use_en_cn_mode:
                                                        text_to_docx.convert_xml_to_docx_English2ChineseMode(structured, out_file)
                                                else:
                                                        text_to_docx.convert_xml_to_docx(structured, out_file)
                                except Exception as e:
                                        errors.append((slide, str(e)))

                # 合并文件写出（如果需要）
                if same_file and combined_structured:
                        out_name = self.filename_tpl.GetValue() or "combined.docx"
                        out_file = os.path.join(out_dir, out_name)
                        try:
                                if text_only:
                                        d = Document()
                                        for textbox in combined_structured:
                                                for line_runs in textbox:
                                                        line_text = ''.join([r.text for r in line_runs if r.text])
                                                        if line_text.strip():
                                                                d.add_paragraph(line_text)
                                        d.save(out_file)
                                else:
                                        if use_en_cn_mode:
                                                text_to_docx.convert_xml_to_docx_English2ChineseMode(combined_structured, out_file)
                                        else:
                                                text_to_docx.convert_xml_to_docx(combined_structured, out_file)
                        except Exception as e:
                                errors.append(("合并", str(e)))

                # 完成后处理
                if errors:
                        err_msg = '\n'.join([f"{s}: {m}" for s,m in errors[:5]]) # 只显示前5个错误以免消息框过大
                        if len(errors) > 5:
                                err_msg += f"\n... 等共 {len(errors)} 个错误"
                        wx.MessageBox(_(u"部分导出失败，请查看日志。") + "\n" + err_msg, _(u"完成(有错误)"), wx.OK|wx.ICON_WARNING)
                else:
                        wx.MessageBox(_(u"导出全部完成。"), _(u"完成"), wx.OK|wx.ICON_INFORMATION)


        def __del__( self ):
                # 避免在对象已被销毁时再次调用 Destroy 导致 RuntimeError
                try:
                        if self and not self.IsDestroyed():
                                self.Destroy()
                except Exception:
                        pass
if __name__ == '__main__':
        app = wx.App(False)
        frame = output(None)
        frame.SetTitle("enbx_to_docx - 导出设置")
        frame.Show()
        app.MainLoop()